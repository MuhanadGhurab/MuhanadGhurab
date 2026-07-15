#!/usr/bin/env python3
"""Deterministic ATS resume builder: PDF + DOCX + preview PNG from resume-data.json."""
from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Twips
from pypdf import PdfReader
from reportlab.lib.colors import Color, HexColor, white
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch, mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

ROOT = Path(__file__).resolve().parents[2]
RESUME_DIR = ROOT / "resume"
DATA_PATH = RESUME_DIR / "source" / "resume-data.json"
PDF_PATH = RESUME_DIR / "Muhanad-Ghurab-ATS-Resume.pdf"
DOCX_PATH = RESUME_DIR / "Muhanad-Ghurab-ATS-Resume.docx"
PREVIEW_PATH = RESUME_DIR / "Muhanad-Ghurab-ATS-Resume-preview.png"

NAVY = HexColor("#0F172A")
BLUE = HexColor("#2563EB")
SLATE = HexColor("#475569")
TEXT = HexColor("#111827")


def load_data() -> dict:
    data = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    validate_data(data)
    return data


def validate_data(data: dict) -> None:
    ident = data["identity"]
    assert ident["email"] == "muhanadghurab@gmail.com"
    assert ident["github_username"] == "MuhanadGhurab"
    assert ident["phone"] == "+966 59 700 4895"
    assert "utm_" not in ident["linkedin_url"]
    assert ident["linkedin_url"].startswith("https://www.linkedin.com/in/")
    assert all(data["development"])
    for phrase in data["forbidden_phrases"]:
        blob = json.dumps(data)
        assert phrase not in blob or phrase in data["forbidden_phrases"]
    # ensure In Progress wording
    assert any("Security+" in d and "In Progress" in d for d in data["development"])
    assert any("PMP" in d and "In Progress" in d for d in data["development"])


def find_font() -> str:
    candidates = [
        Path(r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\calibri.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"),
    ]
    bold_candidates = [
        Path(r"C:\Windows\Fonts\arialbd.ttf"),
        Path(r"C:\Windows\Fonts\calibrib.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        Path("/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"),
    ]
    for regular, bold in zip(candidates, bold_candidates):
        if regular.exists() and bold.exists():
            pdfmetrics.registerFont(TTFont("ResumeSans", str(regular)))
            pdfmetrics.registerFont(TTFont("ResumeSans-Bold", str(bold)))
            return "ResumeSans"
    return "Helvetica"


def wrap(c: canvas.Canvas, text: str, font: str, size: float, max_width: float) -> list[str]:
    words = text.split()
    lines: list[str] = []
    cur = ""
    for w in words:
        trial = f"{cur} {w}".strip()
        if c.stringWidth(trial, font, size) <= max_width:
            cur = trial
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


def draw_link(c: canvas.Canvas, x: float, y: float, text: str, url: str, font: str, size: float, color=BLUE):
    c.setFillColor(color)
    c.setFont(font, size)
    c.drawString(x, y, text)
    w = c.stringWidth(text, font, size)
    c.linkURL(url, (x, y - 2, x + w, y + size), relative=0)


def build_pdf(data: dict, font: str) -> None:
    bold = f"{font}-Bold" if font == "ResumeSans" else "Helvetica-Bold"
    page_w, page_h = A4
    margin = 0.55 * inch
    width = page_w - 2 * margin
    c = canvas.Canvas(str(PDF_PATH), pagesize=A4)
    c.setTitle(data["meta"]["document_title"].replace("—", "-"))
    c.setAuthor(data["meta"]["author"])
    c.setSubject(data["meta"]["subject"])
    c.setKeywords(", ".join(data["meta"]["keywords"]))
    c.setCreator("Muhanad Ghurab ATS Resume Builder")

    y = page_h - margin
    ident = data["identity"]

    # Name
    c.setFillColor(NAVY)
    c.setFont(bold, 22)
    c.drawString(margin, y - 18, ident["name"].upper())
    y -= 34

    c.setFillColor(SLATE)
    c.setFont(font, 11.5)
    c.drawString(margin, y, ident["title"])
    y -= 14

    c.setStrokeColor(BLUE)
    c.setLineWidth(1.4)
    c.line(margin, y, margin + width, y)
    y -= 16

    # Contact (plain text for ATS + URL annotations)
    c.setFillColor(TEXT)
    c.setFont(font, 8.8)
    contact1 = f"{ident['location']}  |  {ident['phone']}  |  {ident['email']}"
    c.drawString(margin, y, contact1)
    # annotate email
    email_x = margin + c.stringWidth(f"{ident['location']}  |  {ident['phone']}  |  ", font, 8.8)
    email_w = c.stringWidth(ident["email"], font, 8.8)
    c.linkURL(f"mailto:{ident['email']}", (email_x, y - 2, email_x + email_w, y + 10), relative=0)
    y -= 12
    contact2 = f"{ident['linkedin_display']}  |  {ident['github_display']}"
    c.drawString(margin, y, contact2)
    li_w = c.stringWidth(ident["linkedin_display"], font, 8.8)
    c.linkURL(ident["linkedin_url"], (margin, y - 2, margin + li_w, y + 10), relative=0)
    gh_x = margin + c.stringWidth(ident["linkedin_display"] + "  |  ", font, 8.8)
    gh_w = c.stringWidth(ident["github_display"], font, 8.8)
    c.linkURL(ident["github_url"], (gh_x, y - 2, gh_x + gh_w, y + 10), relative=0)
    y -= 18

    def section(title: str):
        nonlocal y
        c.setFillColor(NAVY)
        c.setFont(bold, 11)
        c.drawString(margin, y, title.upper())
        y -= 4
        c.setStrokeColor(BLUE)
        c.setLineWidth(0.8)
        c.line(margin, y, margin + width, y)
        y -= 12

    def body_lines(text: str, size=9.5, leading=11.5):
        nonlocal y
        c.setFillColor(TEXT)
        c.setFont(font, size)
        for line in wrap(c, text, font, size, width):
            c.drawString(margin, y, line)
            y -= leading

    def bullet(text: str, size=9.3, leading=11.2):
        nonlocal y
        c.setFillColor(TEXT)
        c.setFont(font, size)
        prefix = "- "
        avail = width - c.stringWidth(prefix, font, size)
        lines = wrap(c, text, font, size, avail)
        for i, line in enumerate(lines):
            c.drawString(margin, y, prefix if i == 0 else "  ")
            c.drawString(margin + c.stringWidth(prefix, font, size), y, line)
            y -= leading

    # Summary
    section("Professional Summary")
    body_lines(data["summary"], size=9.4, leading=11.3)
    y -= 4

    # Experience
    section("Professional Experience")
    for role in data["experience"]:
        c.setFillColor(NAVY)
        c.setFont(bold, 10)
        left = f"{role['title']}  |  {role['organization']}"
        c.drawString(margin, y, left)
        c.setFont(font, 9)
        c.setFillColor(SLATE)
        dates = role["dates"]
        c.drawRightString(margin + width, y, dates)
        y -= 11
        c.setFillColor(SLATE)
        c.setFont(font, 8.8)
        c.drawString(margin, y, role["location"])
        y -= 11
        for b in role["bullets"]:
            bullet(b)
        y -= 3

    # Projects
    section("Selected Technical Projects")
    for proj in data["projects"]:
        c.setFillColor(NAVY)
        c.setFont(bold, 9.6)
        c.drawString(margin, y, proj["name"])
        y -= 11
        if proj.get("achievement"):
            c.setFillColor(SLATE)
            c.setFont(font, 8.8)
            c.drawString(margin, y, proj["achievement"])
            y -= 10
        body_lines(proj["description"], size=9.1, leading=10.8)
        c.setFont(font, 8.5)
        draw_link(c, margin, y, proj["url"].replace("https://", ""), proj["url"], font, 8.5)
        y -= 12

    # Skills
    section("Technical Skills")
    for cat, items in data["skills"].items():
        c.setFillColor(NAVY)
        c.setFont(bold, 9.2)
        label = f"{cat}: "
        c.drawString(margin, y, label)
        c.setFillColor(TEXT)
        c.setFont(font, 9.1)
        x0 = margin + c.stringWidth(label, bold, 9.2)
        rem = width - (x0 - margin)
        lines = wrap(c, items, font, 9.1, rem)
        # first line continues after label
        first = lines[0] if lines else ""
        c.drawString(x0, y, first)
        y -= 11
        for line in lines[1:]:
            c.drawString(margin, y, line)
            y -= 11

    y -= 2
    section("Education and Achievement")
    edu = data["education"]
    c.setFillColor(NAVY)
    c.setFont(bold, 9.5)
    c.drawString(margin, y, edu["degree"])
    y -= 11
    c.setFillColor(TEXT)
    c.setFont(font, 9.2)
    c.drawString(margin, y, f"{edu['institution']}  |  {edu['year']}")
    y -= 11
    c.setFillColor(SLATE)
    c.setFont(font, 9.0)
    c.drawString(margin, y, f"Achievement: {edu['achievement']}")
    y -= 14

    section("Professional Development")
    for d in data["development"]:
        c.setFillColor(TEXT)
        c.setFont(font, 9.3)
        c.drawString(margin, y, d)
        y -= 11

    y -= 2
    section("Languages")
    c.setFillColor(TEXT)
    c.setFont(font, 9.3)
    c.drawString(margin, y, data["languages"])
    y -= 10

    if y < margin - 5:
        raise SystemExit(f"Content overflow risk: remaining y={y}")

    c.showPage()
    c.save()

    reader = PdfReader(str(PDF_PATH))
    if len(reader.pages) != 1:
        raise SystemExit(f"PDF page count must be 1, got {len(reader.pages)}")


def set_run_font(run, name="Arial", size=10, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def build_docx(data: dict) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    core = doc.core_properties
    core.author = "Muhanad Ghurab"
    core.title = data["meta"]["document_title"]
    core.subject = data["meta"]["subject"]

    ident = data["identity"]
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(ident["name"].upper())
    set_run_font(run, size=22, bold=True, color=(15, 23, 42))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(ident["title"])
    set_run_font(run, size=11.5, color=(71, 85, 105))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{ident['location']} | {ident['phone']} | ")
    set_run_font(run, size=9)
    add_hyperlink(p, ident["email"], f"mailto:{ident['email']}")
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    add_hyperlink(p2, ident["linkedin_display"], ident["linkedin_url"])
    run = p2.add_run(" | ")
    set_run_font(run, size=9)
    add_hyperlink(p2, ident["github_display"], ident["github_url"])

    def heading(text: str):
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(6)
        hp.paragraph_format.space_after = Pt(4)
        r = hp.add_run(text.upper())
        set_run_font(r, size=11, bold=True, color=(15, 23, 42))

    def para(text: str, size=9.4):
        bp = doc.add_paragraph()
        bp.paragraph_format.space_after = Pt(2)
        r = bp.add_run(text)
        set_run_font(r, size=size)

    def bullet(text: str):
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(1)
        bp.clear()
        r = bp.add_run(text)
        set_run_font(r, size=9.2)

    heading("Professional Summary")
    para(data["summary"])

    heading("Professional Experience")
    for role in data["experience"]:
        rp = doc.add_paragraph()
        rp.paragraph_format.space_after = Pt(0)
        r = rp.add_run(f"{role['title']} | {role['organization']}")
        set_run_font(r, size=10, bold=True, color=(15, 23, 42))
        r2 = rp.add_run(f"    {role['dates']}")
        set_run_font(r2, size=9, color=(71, 85, 105))
        para(role["location"], size=8.8)
        for b in role["bullets"]:
            bullet(b)

    heading("Selected Technical Projects")
    for proj in data["projects"]:
        para(proj["name"], size=9.6)
        if proj.get("achievement"):
            para(proj["achievement"], size=8.8)
        para(proj["description"], size=9.1)
        lp = doc.add_paragraph()
        lp.paragraph_format.space_after = Pt(4)
        add_hyperlink(lp, proj["url"], proj["url"])

    heading("Technical Skills")
    for cat, items in data["skills"].items():
        bp = doc.add_paragraph()
        bp.paragraph_format.space_after = Pt(1)
        r = bp.add_run(f"{cat}: ")
        set_run_font(r, size=9.2, bold=True, color=(15, 23, 42))
        r2 = bp.add_run(items)
        set_run_font(r2, size=9.1)

    heading("Education and Achievement")
    edu = data["education"]
    para(edu["degree"], size=9.5)
    para(f"{edu['institution']} | {edu['year']}", size=9.2)
    para(f"Achievement: {edu['achievement']}", size=9.0)

    heading("Professional Development")
    for d in data["development"]:
        para(d, size=9.3)

    heading("Languages")
    para(data["languages"], size=9.3)

    doc.save(str(DOCX_PATH))
    # validate openxml zip
    with zipfile.ZipFile(DOCX_PATH) as zf:
        assert "[Content_Types].xml" in zf.namelist()


def render_preview() -> None:
    # Prefer PyMuPDF if available; else reportlab-only path uses pillow via pymupdf
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise SystemExit("PyMuPDF (fitz) required for preview") from e
    doc = fitz.open(str(PDF_PATH))
    assert len(doc) == 1
    page = doc.load_page(0)
    pix = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
    pix.save(str(PREVIEW_PATH))
    doc.close()


def main() -> None:
    data = load_data()
    font = find_font()
    build_pdf(data, font)
    build_docx(data)
    render_preview()
    # extraction sanity
    text = "\n".join(page.extract_text() or "" for page in PdfReader(str(PDF_PATH)).pages)
    required = [
        "Muhanad Ghurab",
        "Cybersecurity",
        "IT Infrastructure",
        "Tekfen Construction",
        "Saudi Aramco",
        "Windows",
        "Network",
        "Security+",
        "PMP",
        "Enterprise Cybersecurity Lab",
        "SecSky",
        "GitHub",
        "Jeddah International College",
        "In Progress",
    ]
    for term in required:
        if term.lower() not in text.lower():
            raise SystemExit(f"Missing extracted term: {term}")
    for bad in data["forbidden_phrases"]:
        if bad.lower() in text.lower():
            raise SystemExit(f"Forbidden phrase present: {bad}")
    print("PDF:", PDF_PATH, PDF_PATH.stat().st_size)
    print("DOCX:", DOCX_PATH, DOCX_PATH.stat().st_size)
    print("PREVIEW:", PREVIEW_PATH, PREVIEW_PATH.stat().st_size)
    print("pages:", len(PdfReader(str(PDF_PATH)).pages))
    print("font:", font)


if __name__ == "__main__":
    main()
